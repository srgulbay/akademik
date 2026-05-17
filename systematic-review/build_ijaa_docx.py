#!/usr/bin/env python3
"""
build_ijaa_docx.py
------------------
Construct the IJAA-SUBMISSION-READY compact Word manuscript
(MANUSCRIPT_IJAA.docx) for
"Quorum Sensing in Acinetobacter baumannii: ... A Systematic Review (2003-2025)"
targeting the International Journal of Antimicrobial Agents (IJAA).

Differs from build_final_docx.py:
  * Input source: MANUSCRIPT_IJAA.md (~7K-word compact variant)
  * Output target: MANUSCRIPT_IJAA.docx
  * Graphical abstract (graphical_abstract.png) inserted immediately after
    the structured Abstract — the visual TL;DR for IJAA submission.
  * Title-page metadata auto-computed for the compact variant
    (target word count <= 7,500; expected ref count 60-80; 3 main tables;
    7 main figures + 1 graphical abstract).

Production-editor build:
  * Title page with italic species names, subtitle, author/affiliation/corresp.
  * Manuscript-info / metadata page (article type, counts, CRediT, etc.)
  * Highlights page (5 bullets, IJAA 85-char cap)
  * Structured abstract (Background/Objective/Methods/Results/Conclusions),
    word-count capped at 300; keywords block.
  * Auto-generated TOC field (F9 in Word populates it).
  * Sans-serif (Calibri) headings + Times New Roman 11pt body, 1.5 spacing.
  * Superscript numeric citations (without brackets).
  * Italic species & gene names; bold/italic markdown preserved.
  * Native Word tables with bold headers, alternating shading, top-rule only.
  * One-per-page figures with caption BELOW; 6.5 in width; centered.
  * Vancouver references with 0.5" hanging indent, 10pt, blue DOI hyperlinks.
  * Header (running title) and footer (page number) on every page.
  * Line numbers in margin for peer review.
  * Bookmarks for each H1 section.
  * Roman numerals for front matter; Arabic from Section 1.
"""

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
ROOT = Path("/home/user/akademik/systematic-review")
MANUSCRIPT_MD = ROOT / "MANUSCRIPT_IJAA.md"
OUT_DOCX = ROOT / "MANUSCRIPT_IJAA.docx"
FIG_DIR = ROOT / "figures"
TBL_DIR = ROOT / "tables"
GRAPHICAL_ABSTRACT = FIG_DIR / "graphical_abstract.png"

FIGURES = {
    1: FIG_DIR / "figure1_prisma_flow.png",
    2: FIG_DIR / "figure2_trends.png",
    "3A": FIG_DIR / "figure3a_qs_circuit.png",
    "3B": FIG_DIR / "figure3b_second_messengers.png",
    4: FIG_DIR / "figure4_interventions.png",
    5: FIG_DIR / "figure5_rob_heatmap.png",
    6: FIG_DIR / "figure6_network.png",
}

FIGURE_CAPTIONS = {
    1: ("Figure 1.", " PRISMA 2020 flow diagram for study selection. "
        "Records identified through PubMed/MEDLINE (n=340), OpenAlex (n=351) "
        "and Crossref (n=3,060) APIs, screened, assessed and included in the "
        "synthesis (n=338 PubMed-anchored unique records)."),
    2: ("Figure 2.", " Publication volume and topical trends in "
        "Acinetobacter baumannii quorum-sensing research, 2003-2025. Annual "
        "record counts and topic-frequency distribution across four eras "
        "(foundational, expansion, clinical-translational, recent)."),
    "3A": ("Figure 3A.", " The canonical AbaI/AbaR quorum-sensing circuit in "
           "Acinetobacter baumannii. AbaI synthesises 3-OH-C12-HSL from "
           "S-adenosyl-methionine and 3-hydroxydodecanoyl-ACP; the "
           "autoinducer diffuses through the membrane, binds AbaR, and the "
           "AbaR-AHL homodimer activates target promoters (lux-box) driving "
           "biofilm, motility, efflux and virulence regulons. Accessory "
           "regulators AbaM, ABUW_1132 (LysR), and CRISPR-Cas provide "
           "feedback layers."),
    "3B": ("Figure 3B.", " Nucleotide second-messenger integration with the "
           "Acinetobacter baumannii quorum-sensing network. The 3',5'-cAMP "
           "(CavA/VfrAb), c-di-GMP (GGDEF/EAL effectors, EF-P), and "
           "(p)ppGpp (RelA) pools converge on abaI/abaR transcription and "
           "biofilm/motility output, integrating cell density with energy "
           "state and host-derived carbon cues."),
    4: ("Figure 4.", " Intervention-class landscape for QS-targeting therapy "
        "against Acinetobacter baumannii. Eight intervention classes (phage, "
        "natural product, peptide, QQ enzyme, nanoparticle, synthetic small "
        "molecule, repurposed drug, antibody/vaccine) ordered by study count "
        "and stratified by maturity (in vitro / in silico / in vivo / "
        "clinical)."),
    5: ("Figure 5.", " Risk-of-bias heatmap across reporting-quality domains "
        "for the included corpus. Rows: study-design strata (in vitro, "
        "animal, in silico, omics, clinical). Columns: domains (CRIS, "
        "SYRCLE, ARRIVE 2.0, MIQE, CHARMS-modified, JBI/NOS items). Cells "
        "colour-coded low / some concerns / high risk."),
    6: ("Figure 6.", " Integrated regulatory network of the A. baumannii "
        "quorum-sensing system, summarising components mapped in the 218 "
        "full-text studies. Nodes: AbaI, AbaR, AbaM, AbiR/AbiS, BfmRS, "
        "AdeRS/AdeABC/AdeFGH, PmrAB, CavA/VfrAb, RelA, ABUW_1132, csu, bap, "
        "pgaABCD, ompA, acinetobactin. Edges: positive (arrow) and negative "
        "(T-bar) regulation; dashed edges indicate inferential links "
        "awaiting structural confirmation."),
}

TABLE_CAPTIONS = {
    1: ("Table 1.", " Characteristics of included studies (n=340). "
        "Distribution by publication period, study design, organism scope, "
        "and topic focus. Topic-focus categories are multi-label; totals "
        "may exceed 100%."),
    2: ("Table 2.", " QSI/QQ intervention classes identified in the "
        "corpus. Counts and representative agents for each of eight "
        "chemical or biological classes spanning phages, natural products, "
        "peptides, enzymes, nanoparticles, synthetic compounds, repurposed "
        "drugs and antibody/vaccine constructs."),
    3: ("Table 3.", " Acinetobacter baumannii quorum-sensing molecular "
        "network: key targets. Gene/protein, functional class, role, "
        "mention frequency in the priority corpus, and current "
        "druggability evidence for each principal QS-network node."),
}

FIGURE_ANCHORS = {
    "### 3.1 Study selection": [1],
    "### 3.2 Characteristics of included studies": [2],
    "#### 3.3.1 The *abaI*/*abaR* axis": ["3A"],
    "#### 3.3.5 Nucleotide second messengers": ["3B"],
    "#### 3.3.7 Open mechanistic questions": [6],
    "### 3.4 Therapeutic Targeting of QS — Intervention Classes": [4],
    "#### 3.9.2 Risk-of-bias patterns": [5],
}

TABLE_ANCHORS = {
    "### 3.2 Characteristics of included studies": [1],
    "### 3.4 Therapeutic Targeting of QS — Intervention Classes": [2],
    "#### 3.3.7 Open mechanistic questions": [3],
}

# Italic auto-detect terms (bacterial nomenclature, gene names, latinisms).
ITALIC_RX = re.compile(
    r"\b(Acinetobacter baumannii|A\. baumannii|Acinetobacter spp\.|"
    r"Pseudomonas aeruginosa|P\. aeruginosa|Staphylococcus aureus|"
    r"S\. aureus|Candida albicans|C\. elegans|Caenorhabditis elegans|"
    r"Galleria mellonella|G\. mellonella|"
    r"abaI|abaR|abaM|abiS|abiR|"
    r"adeABC|adeFGH|adeIJK|adeRS|adeB|"
    r"bfmRS|bfmR|bfmS|pmrAB|"
    r"csuA/B|csuA|csuB|csuC|csuD|csuE|"
    r"pgaABCD|pgaD|bap|ompA|relA|dksA|"
    r"luxI|luxR|lasI|lasR|rhlI|rhlR|pqsR|mvfR|"
    r"aidA|momL|aaL|pvdQ|vfrAb|cavA|aaR4b|"
    r"in vitro|in vivo|in silico|ex vivo|de novo|et al\.|i\.e\.|e\.g\.)\b"
)

CITE_RX = re.compile(r"\[(\d+(?:\s*[,–-]\s*\d+)*(?:\s*,\s*\d+)*)\]")
DOI_RX = re.compile(r"(doi:\s*)(10\.\S+?)(?=[\s)\];,]|$)", re.IGNORECASE)
URL_RX = re.compile(r"(https?://\S+)")


# ---------------------------------------------------------------------------
# Low-level OOXML helpers
# ---------------------------------------------------------------------------

def _oxml(tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), v)
    return el


def _shade(cell, fill_hex):
    """Set background shading on a table cell (e.g. 'F5F5F5')."""
    tcPr = cell._tc.get_or_add_tcPr()
    # Remove existing shd if present (idempotent)
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _cell_top_align(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(old)
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "top")
    tcPr.append(va)


def _cell_borders_top_bottom_only(cell, top=True, bottom=False, weight="6",
                                  color="000000"):
    """Apply only top/bottom borders (no vertical lines)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        if (edge == "top" and top) or (edge == "bottom" and bottom):
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), weight)
            b.set(qn("w:color"), color)
        else:
            b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_bookmark(paragraph, name, bookmark_id):
    """Wrap the paragraph contents with bookmarkStart/bookmarkEnd."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_field(paragraph, instr_text, default_text=""):
    """Insert a Word field (e.g. PAGE, NUMPAGES, TOC, HYPERLINK)."""
    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r.append(fld_begin)

    instr_run = paragraph.add_run()
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr_text
    instr_run._r.append(instr_el)

    sep_run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_sep)

    if default_text:
        paragraph.add_run(default_text)

    end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)
    return run


def add_hyperlink(paragraph, url, text, color_hex="0563C1", underline=True,
                  size_pt=10):
    """Insert a clickable hyperlink run."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    rPr.append(color)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rFonts)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_section_line_numbers(section, count_by=1, distance_twips=200,
                             restart="continuous"):
    """Add continuous line numbering in the margin for peer review."""
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:lnNumType")):
        sectPr.remove(old)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), str(count_by))
    ln.set(qn("w:distance"), str(distance_twips))
    ln.set(qn("w:restart"), restart)
    sectPr.append(ln)


def set_section_pgnum_format(section, fmt="decimal", start=None):
    """Configure section's page-number format (decimal | lowerRoman)."""
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sectPr.append(pg)


def hard_page_break(doc):
    """Insert a true Word page-break run."""
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)
    return p


def configure_styles(doc):
    """
    Set:
      Normal       Times New Roman 11pt, 1.5 spacing, 0pt/6pt before/after
      Heading 1    Calibri 14pt bold (sans-serif)
      Heading 2    Calibri 12pt bold
      Heading 3    Calibri 11pt bold italic
      Heading 4    Calibri 11pt bold
    """
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

    heading_specs = [
        (1, 14, True, False),
        (2, 12, True, False),
        (3, 11, True, True),
        (4, 11, True, False),
    ]
    for level, size, bold, italic in heading_specs:
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        rpr_h = style.element.get_or_add_rPr()
        rfh = rpr_h.find(qn("w:rFonts"))
        if rfh is None:
            rfh = OxmlElement("w:rFonts")
            rpr_h.append(rfh)
        rfh.set(qn("w:ascii"), "Calibri")
        rfh.set(qn("w:hAnsi"), "Calibri")
        rfh.set(qn("w:eastAsia"), "Calibri")
        rfh.set(qn("w:cs"), "Calibri")
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True


def set_all_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


# ---------------------------------------------------------------------------
# Disable revisions / track changes / comments
# ---------------------------------------------------------------------------

def disable_revisions(doc):
    settings_xml = doc.settings.element
    for tag in ("w:trackChanges", "w:revisionView"):
        for old in settings_xml.findall(qn(tag)):
            settings_xml.remove(old)
    # Explicitly hide nothing extra (no revisions to hide).


# ---------------------------------------------------------------------------
# Inline formatting (italic auto, bold, citation superscript)
# ---------------------------------------------------------------------------

def _emit_with_auto_italic(paragraph, content, italic=False, bold=False,
                           size=None):
    """Emit `content` as runs, auto-italicising species/gene tokens."""
    if not content:
        return
    last = 0
    for m in ITALIC_RX.finditer(content):
        if m.start() > last:
            run = paragraph.add_run(content[last:m.start()])
            run.italic = italic
            run.bold = bold
            if size:
                run.font.size = size
        run = paragraph.add_run(m.group(0))
        run.italic = True
        run.bold = bold
        if size:
            run.font.size = size
        last = m.end()
    if last < len(content):
        run = paragraph.add_run(content[last:])
        run.italic = italic
        run.bold = bold
        if size:
            run.font.size = size


def _superscript_citation(paragraph, inner, base_size=None):
    """Render a numeric citation [3,7,15] as the superscript '3,7,15'.

    Strips brackets. Per the production-editor spec, the citation should come
    BEFORE the sentence's terminal period — we do not handle period swap here
    because we only see one bracket-token at a time; the manuscript text is
    already in a "previous studies [3,7]." form which produces "previous
    studies 3,7." with the superscript number placed before the period.
    """
    run = paragraph.add_run(inner)
    run.font.superscript = True
    if base_size:
        run.font.size = base_size


def add_formatted_runs(paragraph, text, base_italic=False, base_bold=False,
                       base_size=None):
    """
    Tokenise a markdown line into bold/italic/citation/text segments and
    append corresponding runs to `paragraph`.

    Handles:
      **bold**          -> bold runs
      *italic*          -> italic runs  (gene names in markdown)
      [3,7,15]          -> superscript runs without brackets
      Plain text        -> auto-italicises species/gene tokens via ITALIC_RX
    """
    tokens = []
    i = 0
    n = len(text)
    while i < n:
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end != -1:
                tokens.append(("bold", text[i + 2:end]))
                i = end + 2
                continue
        if text[i] == "*" and (i == 0 or text[i - 1] != "*"):
            end = text.find("*", i + 1)
            if end != -1 and (end + 1 == n or text[end + 1] != "*"):
                tokens.append(("italic", text[i + 1:end]))
                i = end + 1
                continue
        m = CITE_RX.match(text, i)
        if m:
            tokens.append(("cite", m.group(1)))
            i = m.end()
            continue
        j = i
        while j < n:
            if text.startswith("**", j):
                break
            if text[j] == "*" and (j == 0 or text[j - 1] != "*"):
                nxt = text.find("*", j + 1)
                if nxt != -1 and (nxt + 1 == n or text[nxt + 1] != "*"):
                    break
            if text[j] == "[" and CITE_RX.match(text, j):
                break
            j += 1
        if j == i:
            tokens.append(("text", text[i]))
            i += 1
        else:
            tokens.append(("text", text[i:j]))
            i = j

    for kind, content in tokens:
        if kind == "cite":
            _superscript_citation(paragraph, content, base_size=base_size)
        elif kind == "italic":
            _emit_with_auto_italic(paragraph, content,
                                   italic=True, bold=base_bold, size=base_size)
        elif kind == "bold":
            _emit_with_auto_italic(paragraph, content,
                                   italic=base_italic, bold=True,
                                   size=base_size)
        else:
            _emit_with_auto_italic(paragraph, content,
                                   italic=base_italic, bold=base_bold,
                                   size=base_size)


# ---------------------------------------------------------------------------
# Word-count helpers
# ---------------------------------------------------------------------------

def load_manuscript():
    return MANUSCRIPT_MD.read_text(encoding="utf-8")


def estimate_word_count(md_text):
    """Words in Sections 1-5 prose (excluding refs, suppl, tables)."""
    in_main = False
    in_code = False
    in_table = False
    count = 0
    section_re = re.compile(r"^##\s+(\d+)\.\s")
    for line in md_text.splitlines():
        s = line.strip()
        if s.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        m = section_re.match(s)
        if m:
            sec = int(m.group(1))
            in_main = 1 <= sec <= 5
            continue
        if s.startswith("## References") or s.startswith("## Supplementary"):
            in_main = False
            continue
        if not in_main:
            continue
        # Skip table rows entirely from the main-text word count.
        if s.startswith("|"):
            continue
        cleaned = re.sub(r"[#*_>`]", " ", s)
        cleaned = re.sub(r"\bdoi:\S+", "", cleaned)
        cleaned = re.sub(r"\[\d[\d,\-–\s]*\]", "", cleaned)
        words = [w for w in cleaned.split() if any(ch.isalpha() for ch in w)]
        count += len(words)
    return count


def abstract_word_count():
    total = 0
    for label, body in ABSTRACT_BLOCKS:
        cleaned = re.sub(r"[*_]", "", body)
        cleaned = re.sub(r"\[\d[\d,\-–\s]*\]", "", cleaned)
        total += len([w for w in cleaned.split() if any(c.isalpha() for c in w)])
    return total


# ---------------------------------------------------------------------------
# Native Word table from markdown table block
# ---------------------------------------------------------------------------

def parse_md_table(lines):
    table_lines = [ln for ln in lines
                   if ln.strip().startswith("|") and ln.strip().endswith("|")]
    if not table_lines:
        return None, None
    cleaned = []
    for ln in table_lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", c) for c in cells if c):
            continue
        cleaned.append(cells)
    if not cleaned:
        return None, None
    header, *rows = cleaned
    return header, rows


def add_caption_above(doc, caption_tuple, kind="Table"):
    label, body = caption_tuple
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = "Times New Roman"
    _emit_with_auto_italic(p, body, italic=False, bold=False, size=Pt(10))
    for r in p.runs:
        r.font.size = Pt(10)
    return p


def add_caption_below(doc, caption_tuple, kind="Figure"):
    label, body = caption_tuple
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_with_previous = True
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = "Times New Roman"
    # Caption text portion: italic per spec
    _emit_with_auto_italic(p, body, italic=True, bold=False, size=Pt(10))
    for r in p.runs:
        r.font.size = Pt(10)
    return p


def add_styled_word_table(doc, header, rows, caption=None):
    """Insert a styled native Word table with caption ABOVE it.

    Style: bold header with thin black top+bottom borders, body rows with
    no vertical borders, alternating very-light-gray shading every other
    row, 9pt body / 10pt header, top-aligned cells, first column slightly
    wider for category labels.
    """
    if caption:
        add_caption_above(doc, caption, kind="Table")

    ncols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths
    text_width_in = 6.5
    if ncols == 1:
        widths = [text_width_in]
    else:
        first_col_extra = 0.8 if ncols >= 3 else 0.3
        rest = (text_width_in - (text_width_in / ncols + first_col_extra)) / (ncols - 1)
        first = text_width_in / ncols + first_col_extra
        widths = [first] + [rest] * (ncols - 1)
    for ri in range(len(table.rows)):
        for ci, w in enumerate(widths):
            table.rows[ri].cells[ci].width = Inches(w)

    # Header row
    for ci, htxt in enumerate(header):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        add_formatted_runs(p, htxt, base_bold=True, base_size=Pt(10))
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
        _cell_borders_top_bottom_only(cell, top=True, bottom=True,
                                      weight="8", color="000000")
        _cell_top_align(cell)

    # Body rows
    for ri, row in enumerate(rows, start=1):
        for ci in range(ncols):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            content = row[ci] if ci < len(row) else ""
            add_formatted_runs(p, content, base_size=Pt(9))
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"
            is_last = ri == len(rows)
            _cell_borders_top_bottom_only(cell, top=False,
                                          bottom=is_last,
                                          weight="6", color="000000")
            _cell_top_align(cell)
            if ri % 2 == 0:
                _shade(cell, "F5F5F5")

    # Trailing whitespace paragraph
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(6)


def insert_table_file(doc, table_id):
    file_map = {
        1: TBL_DIR / "table1_characteristics.md",
        2: TBL_DIR / "table2_interventions.md",
        3: TBL_DIR / "table3_qs_targets.md",
    }
    path = file_map.get(table_id)
    if not path or not path.exists():
        p = doc.add_paragraph(f"[Table {table_id} pending]")
        p.runs[0].italic = True
        return
    lines = path.read_text(encoding="utf-8").splitlines()
    header, rows = parse_md_table(lines)
    if header is None:
        p = doc.add_paragraph(f"[Table {table_id} — could not parse]")
        p.runs[0].italic = True
        return
    add_styled_word_table(doc, header, rows, caption=TABLE_CAPTIONS[table_id])


# ---------------------------------------------------------------------------
# Figure insertion (one per page, caption BELOW, 6.5 in)
# ---------------------------------------------------------------------------

def insert_figure(doc, fig_id, embedded_set):
    path = FIGURES.get(fig_id)
    cap = FIGURE_CAPTIONS.get(fig_id, (f"Figure {fig_id}.", ""))
    if not path or not path.exists():
        p = doc.add_paragraph(
            f"[Figure {fig_id} placeholder — image generation pending]")
        p.runs[0].italic = True
        return
    # Page break before each figure to keep it isolated
    hard_page_break(doc)
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        run = p.add_run()
        run.add_picture(str(path), width=Inches(6.5))
        embedded_set.add(fig_id)
    except Exception as e:
        p = doc.add_paragraph(
            f"[Figure {fig_id} placeholder — embed failed: {e}]")
        p.runs[0].italic = True
        return
    add_caption_below(doc, cap, kind="Figure")


# ---------------------------------------------------------------------------
# References handling (with DOI hyperlinking)
# ---------------------------------------------------------------------------

REF_LINE_RX = re.compile(r"^(\d+)\.\s+(.*)$")


def parse_references(md_text):
    refs = []
    inside = False
    for line in md_text.splitlines():
        if line.strip().startswith("## References"):
            inside = True
            continue
        if inside and line.startswith("## "):
            break
        if not inside:
            continue
        m = REF_LINE_RX.match(line.strip())
        if m:
            refs.append((int(m.group(1)), m.group(2).strip()))
    return refs


def _add_reference_run(paragraph, text):
    """Render reference text with italic journal, blue DOI hyperlink."""
    # Replace doi:... with a hyperlink. Split text on DOI tokens.
    last = 0
    for m in DOI_RX.finditer(text):
        before = text[last:m.start()]
        if before:
            add_formatted_runs(paragraph, before, base_size=Pt(10))
        doi_label = m.group(1) + m.group(2)
        url = "https://doi.org/" + m.group(2)
        add_hyperlink(paragraph, url, doi_label, size_pt=10)
        last = m.end()
    if last < len(text):
        add_formatted_runs(paragraph, text[last:], base_size=Pt(10))
    for r in paragraph.runs:
        if r.font.size is None:
            r.font.size = Pt(10)


def add_references_section(doc, refs):
    p = doc.add_heading("References", level=1)
    add_bookmark(p, "sec_references", 7000)
    for n, text in refs:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(1.27)
        pf.first_line_indent = Cm(-1.27)
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.0
        run = p.add_run(f"{n}. ")
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        _add_reference_run(p, text)


# ---------------------------------------------------------------------------
# Title-page / metadata page / highlights / abstract content
# ---------------------------------------------------------------------------

HIGHLIGHTS = [
    "AbaI/AbaR axis integrates BfmRS, AdeRS, PmrAB and second messengers.",
    "Seven QSI/QQ classes mapped; phages lead clinical maturity in A. baumannii.",
    "Biofilm modulation in >75% of corpus; in vitro reductions cluster at 50-80%.",
    "In vivo data are ~15% of studies; ARRIVE 2.0 uneven, PK data scarce.",
    "No QSI past Phase I in A. baumannii; phage compassionate use leads.",
]


def _verify_highlights():
    over = [(i + 1, h, len(h)) for i, h in enumerate(HIGHLIGHTS) if len(h) > 85]
    return over


ABSTRACT_BLOCKS = [
    ("Background:",
     "Acinetobacter baumannii — particularly carbapenem-resistant (CRAB) "
     "and extensively drug-resistant (XDR) phenotypes — is a WHO Priority-1 "
     "pathogen with a constrained antibiotic pipeline. Quorum sensing (QS) "
     "coordinates biofilm formation, virulence-factor expression and "
     "resistance signalling, and is a tractable anti-virulence target."),
    ("Objective:",
     "To characterise (i) the molecular architecture of the A. baumannii QS "
     "network, (ii) QS-targeting interventions and effect sizes, (iii) in "
     "vivo and clinical evidence, and (iv) methodological quality of the "
     "field."),
    ("Methods:",
     "PRISMA 2020-compliant systematic review (PROSPERO pending). "
     "MEDLINE/PubMed was searched (1 Jan 2003 - 31 Dec 2025) and "
     "supplemented through OpenAlex/Crossref APIs; multi-database "
     "integration with Scopus, Web of Science, Embase and Cochrane is "
     "ongoing. Records were de-duplicated by DOI/PMID/title-year. Two "
     "hundred and twenty-eight full-text papers underwent structured "
     "data and claim extraction. Quality was appraised using "
     "design-appropriate frameworks (CRIS, SYRCLE, JBI/Newcastle-Ottawa, "
     "CHARMS-modified, MIQE/MINSEQE)."),
    ("Results:",
     "Three hundred and forty unique records met inclusion criteria (338 "
     "after deduplication). The publication rate has tripled since 2018; "
     "24% of the corpus appeared in 2024-2025. In vitro (25%) and omics "
     "(20%) studies dominate; animal models contribute 15% and clinical "
     "studies 4%. Biofilm modulation (77%), antibiotic-resistance "
     "interaction (76%) and virulence attenuation (61%) lead topical "
     "coverage; the abaI/abaR axis is addressed in 52% of papers. Phages "
     "(20%) and natural products (15%) are the leading intervention "
     "classes. Median in vitro biofilm-mass reductions cluster at 50-80% "
     "at sub-MIC. In vivo evidence comprises ~50 animal studies "
     "(predominantly G. mellonella and murine) with heterogeneous design. "
     "Clinical evidence is limited to observational data and pilot "
     "pharmacokinetic measurements (LC-MS/MS quantification of 3-OH-C12-HSL "
     "in burn-patient plasma). No QSI has progressed beyond Phase I."),
    ("Conclusions:",
     "QS targeting in A. baumannii is mechanistically credible and "
     "supported by accumulating preclinical data, but clinical translation "
     "is bottlenecked by formulation, pharmacokinetic characterisation "
     "and the absence of validated QS biomarkers. Adjunctive "
     "QSI-antibiotic combinations and phage cocktails are the nearest-term "
     "clinical strategies; a unified minimum-reporting dataset would "
     "accelerate cross-study synthesis."),
]

KEYWORDS = ("Acinetobacter baumannii; quorum sensing; quorum quenching; "
            "AbaI; AbaR; biofilm; phage therapy; antimicrobial resistance; "
            "ESKAPE pathogens; systematic review")

SUPPLEMENTARY_ITEMS = [
    ("S1", "PRISMA 2020 checklist (supplementary/S1_prisma2020_checklist.md)."),
    ("S2", "Full search strategies for all databases (02-search-strategies.md)."),
    ("S3", "Multi-database merge log and de-duplication report (literature/merged_unique.csv, literature/prisma_flow_data.json)."),
    ("S4", "Data extraction form (04-data-extraction-form.md)."),
    ("S5", "Categorised corpus (literature/categorized.csv)."),
    ("S6", "Evidence claims database (literature/evidence_claims.json)."),
    ("S7", "Full bibliography of 340 records (bibliography.md)."),
    ("S8", "Citation shortlist - top 60 by importance (literature/citation_shortlist.csv)."),
    ("S9", "Extended Synthesis — long-form (~15K word) variant of this review with full methodological detail and extended discussion (MANUSCRIPT_FINAL.docx)."),
]

CREDIT_ROLES = [
    ("Conceptualization", "[AB, CD]"),
    ("Methodology", "[AB]"),
    ("Data curation", "[AB]"),
    ("Formal analysis", "[AB, CD]"),
    ("Investigation", "[AB, CD]"),
    ("Writing — original draft", "[AB]"),
    ("Writing — review & editing", "[AB, CD]"),
    ("Visualization", "[AB]"),
    ("Supervision", "[CD]"),
]


def build_title_page(doc):
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run("Quorum Sensing in ")
    r1.bold = True
    r1.font.size = Pt(18)
    r1.font.name = "Times New Roman"
    r2 = p.add_run("Acinetobacter baumannii")
    r2.bold = True
    r2.italic = True
    r2.font.size = Pt(18)
    r2.font.name = "Times New Roman"
    r3 = p.add_run(": Molecular Architecture, Therapeutic Targeting and "
                   "Translational Horizons — A Systematic Review (2003–2025)")
    r3.bold = True
    r3.font.size = Pt(18)
    r3.font.name = "Times New Roman"

    # Spacer (>= 1 cm)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(18)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    sub = p.add_run("Submitted to: International Journal of "
                    "Antimicrobial Agents")
    sub.italic = True
    sub.font.size = Pt(11)
    sub.font.name = "Times New Roman"

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    a1 = p.add_run("[Author A]")
    a1.font.size = Pt(11)
    a1.font.name = "Times New Roman"
    sup1 = p.add_run("1")
    sup1.font.size = Pt(11)
    sup1.font.superscript = True
    sep = p.add_run(", ")
    sep.font.size = Pt(11)
    a2 = p.add_run("[Author B]")
    a2.font.size = Pt(11)
    a2.font.name = "Times New Roman"
    sup2 = p.add_run("1,2")
    sup2.font.size = Pt(11)
    sup2.font.superscript = True
    sep = p.add_run(", ")
    sep.font.size = Pt(11)
    a3 = p.add_run("[Author C]")
    a3.font.size = Pt(11)
    a3.font.name = "Times New Roman"
    sup3 = p.add_run("2,*")
    sup3.font.size = Pt(11)
    sup3.font.superscript = True

    # Affiliations
    affiliations = [
        ("1", "[Department of Medical Microbiology, Institution A, "
              "City, Country]"),
        ("2", "[Department of Infectious Diseases, Institution B, "
              "City, Country]"),
    ]
    for sup, txt in affiliations:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        s = p.add_run(sup)
        s.font.size = Pt(10)
        s.font.superscript = True
        body = p.add_run(" " + txt)
        body.italic = True
        body.font.size = Pt(10)
        body.font.name = "Times New Roman"

    # Corresponding author
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    star = p.add_run("*")
    star.font.size = Pt(10)
    star.font.superscript = True
    label = p.add_run(" Corresponding author: ")
    label.bold = True
    label.font.size = Pt(10)
    label.font.name = "Times New Roman"
    val = p.add_run("[Author C], [Postal address]. "
                    "E-mail: corresponding.author@institution.example")
    val.font.size = Pt(10)
    val.font.name = "Times New Roman"


def build_metadata_page(doc, word_count_main, ref_count, fig_count,
                        table_count):
    h = doc.add_heading("Manuscript Information", level=1)

    meta_pairs = [
        ("Article type", "Systematic Review"),
        ("Running head",
         "Quorum sensing in A. baumannii — systematic review"),
        ("Target journal",
         "International Journal of Antimicrobial Agents (IJAA)"),
        ("Word count (main text, Sections 1–5)", f"{word_count_main:,}"),
        ("Abstract word count", f"{abstract_word_count()}"),
        ("Reference count", str(ref_count)),
        ("Number of tables", f"{table_count} (main text)"),
        ("Number of figures", f"{fig_count}"),
        ("Number of supplementary files",
         f"{len(SUPPLEMENTARY_ITEMS)} (S1–S{len(SUPPLEMENTARY_ITEMS)})"),
        ("Highlights", f"{len(HIGHLIGHTS)} bullets, each ≤ 85 characters"),
        ("Keywords", f"{KEYWORDS.count(';') + 1}"),
        ("Protocol registration", "PROSPERO (CRD pending)"),
    ]
    _emit_keyvalue_table(doc, meta_pairs)

    # Acknowledgements / Funding / COI / Data
    _h2(doc, "Acknowledgements")
    _para_italic(doc, "[To be completed.]")

    _h2(doc, "Funding")
    _para_italic(doc, "[To be completed. The authors received no specific "
                      "funding for this work / Funding source to be declared.]")

    _h2(doc, "Conflicts of interest")
    _para_italic(doc, "[To be completed. The authors declare no competing "
                      "interests / Specific conflicts to be declared.]")

    _h2(doc, "Data availability statement")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_formatted_runs(
        p,
        "All data extracted from primary studies are available in "
        "Supplementary Materials S5 and S6. Custom code (data extraction, "
        "deduplication, categorisation, evidence-claim mining, figure "
        "generation) is available in the repository accompanying this "
        "manuscript.")

    _h2(doc, "Author contributions (CRediT)")
    for role, people in CREDIT_ROLES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(role + ": ")
        r.bold = True
        r.font.size = Pt(11)
        v = p.add_run(people)
        v.font.size = Pt(11)


def _h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    return p


def _para_italic(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = True


def _emit_keyvalue_table(doc, pairs):
    """Render label/value list as a borderless two-column table for crisp
    alignment on the metadata page."""
    table = doc.add_table(rows=len(pairs), cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    label_w = Inches(2.4)
    value_w = Inches(4.1)
    for ri, (label, value) in enumerate(pairs):
        for ci, w in [(0, label_w), (1, value_w)]:
            table.rows[ri].cells[ci].width = w
        c1 = table.rows[ri].cells[0]
        c1.text = ""
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(label + ":")
        r1.bold = True
        r1.font.size = Pt(11)
        c2 = table.rows[ri].cells[1]
        c2.text = ""
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_after = Pt(2)
        add_formatted_runs(p2, value, base_size=Pt(11))
        # No borders
        for c in (c1, c2):
            _cell_borders_top_bottom_only(c, top=False, bottom=False)
            _cell_top_align(c)


def build_highlights_page(doc):
    h = doc.add_heading("Highlights", level=1)
    add_bookmark(h, "sec_highlights", 7001)
    over = _verify_highlights()
    for i, hl in enumerate(HIGHLIGHTS):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(6)
        add_formatted_runs(p, hl, base_size=Pt(11))
        for r in p.runs:
            if r.font.size is None:
                r.font.size = Pt(11)
    if over:
        p = doc.add_paragraph()
        r = p.add_run(
            f"[Editorial flag: {len(over)} highlight(s) exceed 85 chars.]")
        r.italic = True
        r.font.size = Pt(9)


def build_abstract_page(doc):
    h = doc.add_heading("Abstract", level=1)
    add_bookmark(h, "sec_abstract", 7002)
    for label, body in ABSTRACT_BLOCKS:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(label + " ")
        r.bold = True
        r.font.size = Pt(11)
        add_formatted_runs(p, body, base_size=Pt(11))
    # Keywords
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Keywords: ")
    r.bold = True
    add_formatted_runs(p, KEYWORDS, base_size=Pt(11))


def build_graphical_abstract_page(doc):
    """Insert the graphical abstract on its own page, immediately after
    the structured abstract. Core IJAA submission element — the visual
    TL;DR of the review. Returns True if the image was embedded, else False."""
    hard_page_break(doc)
    h = doc.add_heading("Graphical Abstract", level=1)
    add_bookmark(h, "sec_graphical_abstract", 7004)

    if not GRAPHICAL_ABSTRACT.exists():
        p = doc.add_paragraph()
        r = p.add_run(
            "[Graphical abstract placeholder — figures/graphical_abstract.png "
            "not yet generated.]")
        r.italic = True
        r.font.size = Pt(10)
        return False

    # Centered image at 6.5 in width
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    try:
        run.add_picture(str(GRAPHICAL_ABSTRACT), width=Inches(6.5))
    except Exception as e:
        err = doc.add_paragraph()
        rerr = err.add_run(f"[Graphical abstract embed failed: {e}]")
        rerr.italic = True
        rerr.font.size = Pt(10)
        return False

    # Caption below
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(12)
    cap.paragraph_format.line_spacing = 1.15
    lab = cap.add_run("Graphical Abstract.")
    lab.bold = True
    lab.font.size = Pt(10)
    lab.font.name = "Times New Roman"
    body = (" Visual TL;DR — quorum sensing (QS) in "
            "Acinetobacter baumannii. The AbaI/AbaR axis, modulated by "
            "BfmRS, AdeRS, PmrAB and nucleotide second messengers "
            "(c-di-GMP, ppGpp, cAMP), drives biofilm formation, motility, "
            "efflux and virulence. Eight intervention classes (phage, "
            "natural product, peptide, QQ enzyme, nanoparticle, synthetic "
            "small molecule, repurposed drug, antibody/vaccine) target the "
            "circuit at distinct nodes. Phages and QSI-antibiotic adjuncts "
            "are the nearest-term clinical strategies; no QSI has yet "
            "progressed beyond Phase I in A. baumannii.")
    _emit_with_auto_italic(cap, body, italic=True, bold=False, size=Pt(10))
    for r in cap.runs:
        r.font.size = Pt(10)
    return True


def build_toc_page(doc):
    h = doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    note = p.add_run(
        "(Right-click below and choose ‘Update Field’, or press F9 in "
        "Word, to populate the table of contents.)")
    note.italic = True
    note.font.size = Pt(9)
    toc_p = doc.add_paragraph()
    # TOC field — levels 1-3, hyperlinks, numbered entries
    add_field(toc_p,
              'TOC \\o "1-3" \\h \\z \\u',
              default_text="Right-click and update field to populate "
                           "the Table of Contents.")


def build_supplementary_section(doc):
    hard_page_break(doc)
    h = doc.add_heading("Supplementary Material", level=1)
    add_bookmark(h, "sec_supplementary", 7003)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    intro = p.add_run("The following supplementary files accompany this "
                      "manuscript:")
    intro.font.size = Pt(11)
    for sid, desc in SUPPLEMENTARY_ITEMS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        r = p.add_run(f"{sid}. ")
        r.bold = True
        add_formatted_runs(p, desc, base_size=Pt(11))


# ---------------------------------------------------------------------------
# Header / Footer (for front-matter and main sections)
# ---------------------------------------------------------------------------

def _add_page_number_footer(section, with_total=False):
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Clear existing runs
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    pre = p.add_run("Page ")
    pre.font.size = Pt(9)
    pre.font.name = "Calibri"
    add_field(p, "PAGE \\* MERGEFORMAT", default_text="1")
    if with_total:
        mid = p.add_run(" of ")
        mid.font.size = Pt(9)
        mid.font.name = "Calibri"
        add_field(p, "NUMPAGES \\* MERGEFORMAT", default_text="1")


def _add_running_header(section, text):
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    run.font.name = "Calibri"


# ---------------------------------------------------------------------------
# Main render of body (Sections 1-5)
# ---------------------------------------------------------------------------

def render_body(doc, md_text, embedded_figs):
    lines = md_text.splitlines()
    start_idx = None
    end_idx = None
    for i, ln in enumerate(lines):
        if re.match(r"^##\s+1\.\s", ln.strip()):
            start_idx = i
            break
    for i, ln in enumerate(lines):
        if ln.strip().startswith("## References"):
            end_idx = i
            break
    if start_idx is None or end_idx is None:
        raise RuntimeError("Couldn't find main text boundaries")

    body_lines = lines[start_idx:end_idx]
    fig_anchors = FIGURE_ANCHORS
    tbl_anchors = TABLE_ANCHORS

    pending = {"figs": [], "tables": []}
    bookmark_counter = [8000]

    def flush_pending():
        for fig_id in pending["figs"]:
            insert_figure(doc, fig_id, embedded_figs)
        for tbl_id in pending["tables"]:
            insert_table_file(doc, tbl_id)
        pending["figs"].clear()
        pending["tables"].clear()

    in_code = False
    table_buffer = []
    last_was_heading = False

    def emit_table_buffer():
        nonlocal table_buffer
        if not table_buffer:
            return
        header, rows = parse_md_table(table_buffer)
        if header is not None:
            add_styled_word_table(doc, header, rows, caption=None)
        table_buffer = []

    i = 0
    n = len(body_lines)
    while i < n:
        line = body_lines[i]
        stripped = line.rstrip()

        if stripped.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            i += 1
            continue

        m1 = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m1:
            emit_table_buffer()
            flush_pending()
            level_hashes = m1.group(1)
            heading_text = m1.group(2).strip()
            level = len(level_hashes) - 1
            level = max(1, min(level, 4))
            if re.match(r"^Table\s+\d", heading_text) or re.match(
                    r"^Figure\s+\d", heading_text):
                i += 1
                continue

            full_heading_line = stripped
            for anchor_key, fig_list in fig_anchors.items():
                if full_heading_line == anchor_key:
                    pending["figs"].extend(fig_list)
            for anchor_key, tbl_list in tbl_anchors.items():
                if full_heading_line == anchor_key:
                    pending["tables"].extend(tbl_list)

            word_level = level
            h = doc.add_heading(heading_text, level=word_level)
            # Add bookmark for top-level (## -> H1) sections
            if word_level == 1:
                slug = re.sub(r"[^A-Za-z0-9]+", "_", heading_text)[:48].strip("_")
                bookmark_counter[0] += 1
                add_bookmark(h, f"sec_{slug}", bookmark_counter[0])
            last_was_heading = True
            i += 1
            continue

        if stripped.startswith("|"):
            table_buffer.append(stripped)
            i += 1
            continue
        else:
            if table_buffer:
                emit_table_buffer()

        if not stripped.strip():
            i += 1
            continue

        if stripped.strip() == "---":
            i += 1
            continue

        bullet_m = re.match(r"^([\-\*])\s+(.*)$", stripped)
        if bullet_m:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_formatted_runs(p, bullet_m.group(2))
            last_was_heading = False
            i += 1
            continue

        num_m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if num_m and len(num_m.group(1)) <= 2:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            add_formatted_runs(p, num_m.group(2))
            last_was_heading = False
            i += 1
            continue

        # Body paragraph
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.5
        # First-line indent except immediately after a heading
        if not last_was_heading:
            pf.first_line_indent = Cm(0.5)
        add_formatted_runs(p, stripped)
        last_was_heading = False
        i += 1

    emit_table_buffer()
    flush_pending()


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------

def main():
    md = load_manuscript()
    word_count_main = estimate_word_count(md)
    refs = parse_references(md)
    ref_count = len(refs)

    doc = Document()
    set_all_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0)
    configure_styles(doc)
    disable_revisions(doc)

    running_title = "Quorum sensing in A. baumannii — systematic review"

    # ---- Section 1 (front matter): Roman numerals page numbering ----
    s0 = doc.sections[0]
    _add_running_header(s0, running_title)
    _add_page_number_footer(s0, with_total=False)
    set_section_pgnum_format(s0, fmt="lowerRoman", start=1)
    set_section_line_numbers(s0, count_by=1, distance_twips=200,
                             restart="newSection")

    # Title page
    build_title_page(doc)
    hard_page_break(doc)

    # Metadata page: 7 main figures + 1 graphical abstract = 8 figures total
    build_metadata_page(doc, word_count_main, ref_count, fig_count=8,
                        table_count=3)
    hard_page_break(doc)

    # TOC
    build_toc_page(doc)
    hard_page_break(doc)

    # Highlights
    build_highlights_page(doc)
    hard_page_break(doc)

    # Abstract
    build_abstract_page(doc)

    # Graphical abstract — the visual TL;DR for IJAA, immediately after abstract
    graphical_abstract_inserted = build_graphical_abstract_page(doc)

    # ---- Section break: switch to Arabic page numbers for main text ----
    main_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    main_section.top_margin = Inches(1.0)
    main_section.bottom_margin = Inches(1.0)
    main_section.left_margin = Inches(1.0)
    main_section.right_margin = Inches(1.0)
    main_section.header.is_linked_to_previous = False
    main_section.footer.is_linked_to_previous = False
    _add_running_header(main_section, running_title)
    _add_page_number_footer(main_section, with_total=True)
    set_section_pgnum_format(main_section, fmt="decimal", start=1)
    set_section_line_numbers(main_section, count_by=1, distance_twips=200,
                             restart="continuous")

    embedded_figs = set()
    render_body(doc, md, embedded_figs)

    # References
    hard_page_break(doc)
    add_references_section(doc, refs)

    # Supplementary
    build_supplementary_section(doc)

    doc.save(str(OUT_DOCX))

    size = OUT_DOCX.stat().st_size
    print(f"WROTE {OUT_DOCX}")
    print(f"FILE_SIZE_BYTES={size}")
    print(f"FILE_SIZE_MB={size/1_048_576:.2f}")
    print(f"WORD_COUNT_MAIN_TEXT={word_count_main}")
    print(f"WORD_COUNT_HARD_CAP=7500 (target <=7000)")
    if word_count_main > 7500:
        print(f"WARNING: word count {word_count_main} > 7500 hard cap")
    elif word_count_main > 7000:
        print(f"NOTE: word count {word_count_main} in safety buffer "
              f"(>7000 target, <=7500 hard cap)")
    print(f"ABSTRACT_WORD_COUNT={abstract_word_count()}")
    print(f"GRAPHICAL_ABSTRACT_INSERTED={graphical_abstract_inserted}")
    print(f"FIGURES_EMBEDDED={sorted(str(k) for k in embedded_figs)}")
    print(f"FIGURES_AVAILABLE={[k for k,v in FIGURES.items() if v.exists()]}")
    print(f"REFERENCES_PARSED={ref_count}")
    over = _verify_highlights()
    if over:
        print(f"HIGHLIGHTS_OVER_85CHAR={over}")
    else:
        print("HIGHLIGHTS_OK=all <= 85 chars")


if __name__ == "__main__":
    main()
