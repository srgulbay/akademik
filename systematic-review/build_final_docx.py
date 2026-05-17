#!/usr/bin/env python3
"""
build_final_docx.py
-------------------
Construct a submission-ready Word manuscript (MANUSCRIPT_FINAL.docx) for
"Quorum Sensing in Acinetobacter baumannii: ... A Systematic Review (2003-2025)"
targeting the International Journal of Antimicrobial Agents.

Parses MANUSCRIPT.md, embeds figures (figures/), renders tables/*.md as native
Word tables, applies italic to bacterial nomenclature and gene names, renders
numeric citations as superscript, and emits a styled, paginated document with
page numbers, references and supplementary sections.
"""

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
ROOT = Path("/home/user/akademik/systematic-review")
MANUSCRIPT_MD = ROOT / "MANUSCRIPT.md"
OUT_DOCX = ROOT / "MANUSCRIPT_FINAL.docx"
FIG_DIR = ROOT / "figures"
TBL_DIR = ROOT / "tables"

FIGURES = {
    1: FIG_DIR / "figure1_prisma_flow.png",
    2: FIG_DIR / "figure2_trends.png",
    "3A": FIG_DIR / "figure3a_qs_circuit.png",
    "3B": FIG_DIR / "figure3b_second_messengers.png",
    4: FIG_DIR / "figure4_interventions.png",
    5: FIG_DIR / "figure5_rob_heatmap.png",
    6: FIG_DIR / "figure6_network.png",
}

FIGURE_CAPTIONS = {
    1: ("Figure 1. PRISMA 2020 flow diagram for study selection. Records "
        "identified through PubMed/MEDLINE (n=340), OpenAlex (n=351) and "
        "Crossref (n=3,060) APIs, screened, assessed and included in the v1 "
        "synthesis (n=338 PubMed-anchored unique records)."),
    2: ("Figure 2. Publication volume and topical trends in Acinetobacter "
        "baumannii quorum-sensing research, 2003-2025. Annual record counts "
        "and topic-frequency distribution across four eras (foundational, "
        "expansion, clinical-translational, recent)."),
    "3A": ("Figure 3A. The canonical AbaI/AbaR quorum-sensing circuit in "
           "Acinetobacter baumannii. AbaI synthesises 3-OH-C12-HSL from "
           "S-adenosyl-methionine and 3-hydroxydodecanoyl-ACP; the autoinducer "
           "diffuses through the membrane, binds AbaR, and the AbaR-AHL "
           "homodimer activates target promoters (lux-box) driving biofilm, "
           "motility, efflux and virulence regulons. Accessory regulators "
           "AbaM, ABUW_1132 (LysR), and CRISPR-Cas provide feedback layers."),
    "3B": ("Figure 3B. Nucleotide second-messenger integration with the "
           "Acinetobacter baumannii quorum-sensing network. The 3',5'-cAMP "
           "(CavA/VfrAb), c-di-GMP (GGDEF/EAL effectors, EF-P), and "
           "(p)ppGpp (RelA) pools converge on abaI/abaR transcription and "
           "biofilm/motility output, integrating cell density with energy "
           "state and host-derived carbon cues."),
    4: ("Figure 4. Intervention-class landscape for QS-targeting therapy "
        "against Acinetobacter baumannii. Eight intervention classes "
        "(phage, natural product, peptide, QQ enzyme, nanoparticle, "
        "synthetic small molecule, repurposed drug, antibody/vaccine) "
        "ordered by study count and stratified by maturity (in vitro / "
        "in silico / in vivo / clinical)."),
    5: ("Figure 5. Risk-of-bias heatmap across reporting-quality domains "
        "for the included corpus. Rows: study-design strata (in vitro, "
        "animal, in silico, omics, clinical). Columns: domains (CRIS, "
        "SYRCLE, ARRIVE 2.0, MIQE, CHARMS-modified, JBI/NOS items). "
        "Cells colour-coded low / some concerns / high risk."),
    6: ("Figure 6. Integrated regulatory network of the A. baumannii "
        "quorum-sensing system, summarising components mapped in the 218 "
        "full-text studies. Nodes: AbaI, AbaR, AbaM, AbiR/AbiS, BfmRS, "
        "AdeRS/AdeABC/AdeFGH, PmrAB, CavA/VfrAb, RelA, ABUW_1132, csu, "
        "bap, pgaABCD, ompA, acinetobactin. Edges: positive (arrow) and "
        "negative (T-bar) regulation; dashed edges indicate inferential "
        "links awaiting structural confirmation."),
}

TABLE_CAPTIONS = {
    1: ("Table 1. Characteristics of included studies (n=340). Distribution "
        "by publication period, study design, organism scope, and topic "
        "focus. Topic-focus categories are multi-label; totals may exceed "
        "100%."),
    2: ("Table 2. QSI/QQ intervention classes identified in the corpus. "
        "Counts and representative agents for each of eight chemical or "
        "biological classes spanning phages, natural products, peptides, "
        "enzymes, nanoparticles, synthetic compounds, repurposed drugs "
        "and antibody/vaccine constructs."),
    3: ("Table 3. Acinetobacter baumannii quorum-sensing molecular network: "
        "key targets. Gene/protein, functional class, role, mention "
        "frequency in the priority corpus, and current druggability "
        "evidence for each principal QS-network node."),
}

# Figures embedded after the line containing each anchor heading;
# value is list of figure keys to insert at the END of that subsection.
FIGURE_ANCHORS = {
    # After 3.1 study selection
    "### 3.1 Study selection": [1],
    "### 3.2 Characteristics of included studies": [2],
    # After 3.3 paragraph mentioning AbaI/AbaR – use 3.3.1 end
    "#### 3.3.1 The abaI/abaR axis: the canonical LuxI/LuxR-type system": ["3A"],
    "#### 3.3.5 Nucleotide second messengers (c-di-GMP, ppGpp, 3′,5′-cAMP) as QS modulators": ["3B"],
    "#### 3.3.7 Open mechanistic questions and unmapped regulons": [6],
    "### 3.4 Therapeutic Targeting of QS — Intervention Classes": [4],
    "#### 3.9.2 Risk of bias assessment summary": [5],
}

# Table anchors — insert Word table at end of named subsection
TABLE_ANCHORS = {
    "### 3.2 Characteristics of included studies": [1],
    "### 3.4 Therapeutic Targeting of QS — Intervention Classes": [2],
    # Table 3 sits inside 3.3 — anchor on 3.3.7 end
    "#### 3.3.7 Open mechanistic questions and unmapped regulons": [3],
}

# Italic gene/protein/species patterns (matched outside markdown asterisks).
ITALIC_RX = re.compile(
    r"\b(Acinetobacter baumannii|A\. baumannii|Acinetobacter spp\.|"
    r"Pseudomonas aeruginosa|P\. aeruginosa|Staphylococcus aureus|"
    r"S\. aureus|Candida albicans|C\. elegans|Caenorhabditis elegans|"
    r"Galleria mellonella|G\. mellonella|abaI|abaR|abaM|abiS|abiR|"
    r"adeABC|adeFGH|adeIJK|adeRS|adeB|bfmRS|bfmR|bfmS|pmrAB|csuA/B|"
    r"csuA|csuB|csuC|csuD|csuE|csuE|pgaABCD|pgaD|bap|ompA|relA|dksA|"
    r"luxI|luxR|lasI|lasR|rhlI|rhlR|pqsR|mvfR|aidA|momL|aaL|pvdQ|"
    r"vfrAb|cavA|aaR4b|N-acyl|in vitro|in vivo|in silico|ex vivo|"
    r"de novo|et al\.|i\.e\.|e\.g\.)\b"
)

CITE_RX = re.compile(r"\[(\d+(?:\s*[,–-]\s*\d+)*(?:\s*,\s*\d+)*)\]")
INLINE_ITALIC_RX = re.compile(r"\*([^*\n]+?)\*")
INLINE_BOLD_RX = re.compile(r"\*\*([^*\n]+?)\*\*")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_border(cell, **kwargs):
    """Add simple borders to a single table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_page_number_footer(section):
    """Inject a PAGE field into the section footer (bottom right)."""
    footer = section.footer
    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* MERGEFORMAT"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break()
    r._r.append(_page_break_xml())


def _page_break_xml():
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    return br


def hard_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)


def configure_styles(doc):
    """Set Normal & heading defaults to TNR 12pt, 1.5 line spacing."""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    # East-Asian fallback
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")

    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)

    for level, size, bold in [(1, 14, True), (2, 12, True), (3, 12, True)]:
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15


def set_margins(doc, inches=1.0):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


# ---------------------------------------------------------------------------
# Inline formatting (italic, bold, superscript citations)
# ---------------------------------------------------------------------------

def add_formatted_runs(paragraph, text, base_italic=False, base_bold=False,
                       base_size=None):
    """
    Append runs to `paragraph` representing `text`, honouring markdown
    *italic* and **bold** markers, rendering numeric citations [1,2] in
    superscript, and italicising bacterial nomenclature/gene tokens.

    Strategy: tokenise the line into segments.
    """
    # First, split by **bold** then *italic* then citations.
    # Approach: walk character-by-character with a small state machine.

    # Tokenise into list of (kind, content)
    tokens = []
    i = 0
    n = len(text)
    while i < n:
        # Bold **...**
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end != -1:
                tokens.append(("bold", text[i + 2:end]))
                i = end + 2
                continue
        # Italic *...*
        if text[i] == "*" and (i == 0 or text[i - 1] != "*"):
            end = text.find("*", i + 1)
            # Avoid matching ** inadvertently
            if end != -1 and (end + 1 == n or text[end + 1] != "*"):
                tokens.append(("italic", text[i + 1:end]))
                i = end + 1
                continue
        # Citation [1,2,3] (numeric only)
        m = CITE_RX.match(text, i)
        if m:
            tokens.append(("cite", m.group(1)))
            i = m.end()
            continue
        # Plain char until next special
        j = i
        while j < n:
            if text.startswith("**", j):
                break
            if text[j] == "*" and (j == 0 or text[j - 1] != "*"):
                # potential italic open
                nxt = text.find("*", j + 1)
                if nxt != -1 and (nxt + 1 == n or text[nxt + 1] != "*"):
                    break
            if text[j] == "[" and CITE_RX.match(text, j):
                break
            j += 1
        if j == i:
            tokens.append(("text", text[i]))
            i += 1
        else:
            tokens.append(("text", text[i:j]))
            i = j

    for kind, content in tokens:
        if kind == "cite":
            # Add "[" + superscript number(s) + "]"  -> spec says superscript;
            # we render whole "[1,2]" as superscript without brackets for
            # journal style.
            run = paragraph.add_run(content)
            run.font.superscript = True
            run.italic = base_italic
            run.bold = base_bold
            if base_size:
                run.font.size = base_size
        elif kind == "italic":
            # Recursively format italic content for nested gene/species
            _emit_italicised(paragraph, content,
                             italic=True, bold=base_bold, size=base_size)
        elif kind == "bold":
            _emit_italicised(paragraph, content,
                             italic=base_italic, bold=True, size=base_size)
        else:
            _emit_italicised(paragraph, content,
                             italic=base_italic, bold=base_bold,
                             size=base_size)


def _emit_italicised(paragraph, content, italic=False, bold=False, size=None):
    """
    Emit plain text, italicising substrings that match ITALIC_RX (bacterial
    species, gene names). Useful for non-markdown-asterisk italic terms.
    """
    if not content:
        return
    last = 0
    for m in ITALIC_RX.finditer(content):
        if m.start() > last:
            run = paragraph.add_run(content[last:m.start()])
            run.italic = italic
            run.bold = bold
            if size:
                run.font.size = size
        run = paragraph.add_run(m.group(0))
        run.italic = True
        run.bold = bold
        if size:
            run.font.size = size
        last = m.end()
    if last < len(content):
        run = paragraph.add_run(content[last:])
        run.italic = italic
        run.bold = bold
        if size:
            run.font.size = size


# ---------------------------------------------------------------------------
# Manuscript parsing
# ---------------------------------------------------------------------------

def load_manuscript():
    return MANUSCRIPT_MD.read_text(encoding="utf-8")


def split_sections(md_text):
    """
    Return list of (heading_or_None, line_or_block_lines).
    We just return the raw lines; section logic is done in render.
    """
    return md_text.splitlines()


# ---------------------------------------------------------------------------
# Word-table rendering from markdown table block
# ---------------------------------------------------------------------------

def parse_md_table(lines):
    """
    Parse markdown table lines into header + rows.
    Skips alignment separator. Returns (header:list[str], rows:list[list[str]]).
    """
    table_lines = [ln for ln in lines
                   if ln.strip().startswith("|") and ln.strip().endswith("|")]
    if not table_lines:
        return None, None
    # Drop the separator row (---|---|...)
    cleaned = []
    for ln in table_lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", c) for c in cells if c):
            continue
        cleaned.append(cells)
    if not cleaned:
        return None, None
    header, *rows = cleaned
    return header, rows


def add_word_table(doc, header, rows, caption=None):
    """Insert a styled native Word table after an optional caption."""
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.paragraph_format.space_before = Pt(6)
        cap_p.paragraph_format.space_after = Pt(3)
        cap_run = cap_p.add_run(caption)
        cap_run.bold = True
        cap_run.italic = True
        cap_run.font.size = Pt(10)

    ncols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    table.autofit = True

    # Header row
    for ci, htxt in enumerate(header):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        add_formatted_runs(p, htxt, base_bold=True, base_size=Pt(10))
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
        set_cell_border(cell)

    for ri, row in enumerate(rows, start=1):
        for ci in range(ncols):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            content = row[ci] if ci < len(row) else ""
            add_formatted_runs(p, content, base_size=Pt(10))
            for run in p.runs:
                run.font.size = Pt(10)
            set_cell_border(cell)

    # Spacing after
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def insert_table_file(doc, table_id):
    """Load tables/tableN_*.md and emit Word table with caption."""
    file_map = {
        1: TBL_DIR / "table1_characteristics.md",
        2: TBL_DIR / "table2_interventions.md",
        3: TBL_DIR / "table3_qs_targets.md",
    }
    path = file_map.get(table_id)
    if not path or not path.exists():
        p = doc.add_paragraph(f"[Table {table_id} pending]")
        p.runs[0].italic = True
        return
    lines = path.read_text(encoding="utf-8").splitlines()
    header, rows = parse_md_table(lines)
    if header is None:
        p = doc.add_paragraph(f"[Table {table_id} — could not parse]")
        p.runs[0].italic = True
        return
    add_word_table(doc, header, rows, caption=TABLE_CAPTIONS[table_id])


# ---------------------------------------------------------------------------
# Figure insertion
# ---------------------------------------------------------------------------

def insert_figure(doc, fig_id, embedded_set):
    path = FIGURES.get(fig_id)
    cap = FIGURE_CAPTIONS.get(fig_id, f"Figure {fig_id}.")
    if not path or not path.exists():
        p = doc.add_paragraph(
            f"[Figure {fig_id} placeholder — image generation pending]")
        p.runs[0].italic = True
        return
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(6.0))
        embedded_set.add(fig_id)
    except Exception as e:
        p = doc.add_paragraph(
            f"[Figure {fig_id} placeholder — embed failed: {e}]")
        p.runs[0].italic = True
        return
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap_p.paragraph_format.space_before = Pt(3)
    cap_p.paragraph_format.space_after = Pt(12)
    cap_run = cap_p.add_run(cap)
    cap_run.italic = True
    cap_run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# References handling
# ---------------------------------------------------------------------------

REF_LINE_RX = re.compile(r"^(\d+)\.\s+(.*)$")


def parse_references(md_text):
    """Return list of (n, raw_text) from the References section."""
    refs = []
    inside = False
    for line in md_text.splitlines():
        if line.strip().startswith("## References"):
            inside = True
            continue
        if inside and line.startswith("## "):
            break
        if not inside:
            continue
        m = REF_LINE_RX.match(line.strip())
        if m:
            refs.append((int(m.group(1)), m.group(2).strip()))
    return refs


def add_references_section(doc, refs):
    h = doc.add_heading("References", level=1)
    for n, text in refs:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.0
        run = p.add_run(f"{n}. ")
        run.font.size = Pt(10)
        run.bold = False
        # Render rest with markdown italic support
        add_formatted_runs(p, text, base_size=Pt(10))
        for r in p.runs:
            r.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Word count
# ---------------------------------------------------------------------------

def estimate_word_count(md_text):
    """
    Words in the main text (Sections 1-5 prose) — excluding the references
    list and the supplementary-material section. Markdown tables, code
    blocks (ASCII art) and bibliographic doi tokens are excluded.
    """
    in_main = False
    in_code = False
    count = 0
    section_re = re.compile(r"^##\s+(\d+)\.\s")
    for line in md_text.splitlines():
        stripped = line.rstrip()
        if stripped.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        s = stripped.strip()
        m = section_re.match(s)
        if m:
            sec = int(m.group(1))
            in_main = 1 <= sec <= 5
            continue
        if s.startswith("## References") or s.startswith("## Supplementary"):
            in_main = False
            continue
        if not in_main:
            continue
        if s.startswith("|"):
            # Count table cell text minimally (table contributes to length)
            cells = [c.strip() for c in s.strip("|").split("|")]
            for c in cells:
                cleaned = re.sub(r"[*_`]", " ", c)
                count += len([w for w in cleaned.split()
                              if any(ch.isalpha() for ch in w)])
            continue
        cleaned = re.sub(r"[#*_>`]", " ", s)
        cleaned = re.sub(r"\bdoi:\S+", "", cleaned)
        words = [w for w in cleaned.split() if any(ch.isalpha() for ch in w)]
        count += len(words)
    return count


# ---------------------------------------------------------------------------
# Title page / metadata blocks
# ---------------------------------------------------------------------------

TITLE = ("Quorum Sensing in Acinetobacter baumannii: Molecular Architecture, "
         "Therapeutic Targeting and Translational Horizons — A Systematic "
         "Review (2003–2025)")

HIGHLIGHTS = [
    ("The Acinetobacter baumannii quorum-sensing (QS) network — anchored by "
     "the AbaI/AbaR LuxI/LuxR-type pair and integrated with the BfmRS, AdeRS "
     "and PmrAB two-component systems plus c-di-GMP, (p)ppGpp and 3′,5′-cAMP "
     "nucleotide second messengers — has been characterised across 22 years "
     "and >300 primary studies."),
    ("Anti-virulence interventions span seven classes (natural products, "
     "synthetic small molecules, phages, nanoparticles, peptides, QQ "
     "enzymes, repurposed drugs), with phage cocktails as the most "
     "clinically advanced modality."),
    ("Biofilm modulation is the most consistently reported outcome (>75% "
     "of corpus), with median in vitro biofilm-mass reductions in the "
     "50–80% range for representative QSIs at sub-MIC concentrations."),
    ("In vivo efficacy data remain modest in volume (~15% of studies) and "
     "heterogeneous in reporting; ARRIVE 2.0 adherence is uneven and "
     "clinical pharmacokinetic data are largely absent."),
    ("LC-MS/MS detection of 3-OH-C12-HSL in burn-patient plasma and phage "
     "compassionate-use protocols represent the leading translational "
     "signals; no QSI has progressed beyond Phase I for A. baumannii."),
]

ABSTRACT_BLOCKS = [
    ("Background.",
     "Acinetobacter baumannii — particularly its carbapenem-resistant "
     "(CRAB) and extensively drug-resistant (XDR) phenotypes — is a WHO "
     "Priority-1 critical pathogen with a constrained antibiotic pipeline. "
     "Quorum sensing (QS) coordinates biofilm formation, virulence-factor "
     "expression and resistance signalling, and has emerged as a tractable "
     "anti-virulence target."),
    ("Objective.",
     "To systematically characterise (i) the molecular architecture of the "
     "A. baumannii QS network, (ii) the spectrum of QS-targeting "
     "interventions and their effect sizes, (iii) the state of in vivo and "
     "clinical evidence, and (iv) the methodological quality of the field."),
    ("Methods.",
     "PRISMA 2020-compliant systematic review (protocol PROSPERO, pending). "
     "MEDLINE/PubMed was searched on 17 May 2026 using a structured Boolean "
     "query (1 Jan 2003 – 31 Dec 2025). The first-pass synthesis is being "
     "extended to Scopus, Web of Science, Embase and Cochrane through "
     "parallel OpenAlex/Crossref API queries (Supplementary S2). Records "
     "were de-duplicated by DOI/PMID/title-year fingerprint. Studies were "
     "categorised by design and intervention class. Evidence claims, "
     "regulator/gene mentions and quantitative outcomes were extracted from "
     "228 full-text papers. Quality was appraised against tool-appropriate "
     "frameworks (CRIS, SYRCLE, JBI/Newcastle-Ottawa, CHARMS-modified, "
     "MIQE/MINSEQE)."),
    ("Results.",
     "Three hundred and forty unique records met initial inclusion criteria "
     "(338 after deduplication). Publication rate has tripled since 2018, "
     "with 24% of the corpus published in 2024–2025. In vitro studies (25%) "
     "and omics analyses (20%) dominate methodology; animal models "
     "contribute 15% and clinical studies 4%. Biofilm modulation (77%), "
     "antibiotic-resistance interaction (76%) and virulence attenuation "
     "(61%) are the most-covered topics; the abaI/abaR axis is the central "
     "regulatory module addressed in 52% of papers. Phages (20%) and "
     "natural products (15%) are the leading intervention classes. Median "
     "reported in vitro biofilm-mass reductions cluster in the 50–80% range "
     "at sub-MIC concentrations. In vivo evidence comprises ~50 animal "
     "studies — predominantly Galleria mellonella and murine — with "
     "heterogeneous design and incomplete ARRIVE 2.0 adherence. Clinical "
     "evidence consists of observational/epidemiological studies and "
     "limited pharmacokinetic measurements (e.g., LC-MS/MS quantification "
     "of 3-OH-C12-HSL in burn-patient plasma). No QSI has progressed "
     "beyond Phase I for A. baumannii."),
    ("Conclusions.",
     "QS targeting in A. baumannii is mechanistically credible and "
     "supported by accumulating preclinical data, but clinical translation "
     "is bottlenecked by formulation, pharmacokinetic characterisation and "
     "the absence of validated QS biomarkers in patient cohorts. "
     "Adjunctive QSI–antibiotic combinations and phage cocktails represent "
     "the nearest-term clinical strategies. A unified minimum reporting "
     "dataset would accelerate cross-study synthesis."),
]

KEYWORDS = ("Acinetobacter baumannii; quorum sensing; quorum quenching; "
            "AbaI; AbaR; biofilm; phage therapy; antimicrobial resistance; "
            "ESKAPE pathogens; systematic review.")

SUPPLEMENTARY_ITEMS = [
    ("S1", "PRISMA 2020 checklist (supplementary/S1_prisma2020_checklist.md)."),
    ("S2", "Full search strategies for all databases (02-search-strategies.md)."),
    ("S3", "Multi-database merge log and de-duplication report (literature/merged_unique.csv, literature/prisma_flow_data.json)."),
    ("S4", "Data extraction form (04-data-extraction-form.md)."),
    ("S5", "Categorised corpus (literature/categorized.csv)."),
    ("S6", "Evidence claims database (literature/evidence_claims.json)."),
    ("S7", "Full bibliography of 340 records (bibliography.md)."),
    ("S8", "Citation shortlist — top 60 by importance (literature/citation_shortlist.csv)."),
]


# ---------------------------------------------------------------------------
# Title-page builder
# ---------------------------------------------------------------------------

def build_title_page(doc, word_count):
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(18)
    title_text = ("Quorum Sensing in ")
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(16)
    run2 = p.add_run("Acinetobacter baumannii")
    run2.bold = True
    run2.italic = True
    run2.font.size = Pt(16)
    run3 = p.add_run(": Molecular Architecture, Therapeutic Targeting and "
                     "Translational Horizons — A Systematic Review (2003–2025)")
    run3.bold = True
    run3.font.size = Pt(16)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Submitted to: International Journal of Antimicrobial Agents")
    run.italic = True
    run.font.size = Pt(12)

    # Authors / affiliations / corresponding author
    for line in [
        "Authors: [To be completed]",
        "Affiliations: [To be completed]",
        "Corresponding author: [To be completed; email]",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Metadata block
    # Provide both Sections 1-5 estimate and the full-document figure for
    # editorial transparency. The "main text" figure here is the conservative
    # Sections 1-5 prose count (excluding references and supplementary).
    meta_lines = [
        ("Article type:", "Systematic Review"),
        ("Word count (main text, Sections 1–5):", f"{word_count:,}"),
        ("Reference count:", "93"),
        ("Tables:", "3 (in main text)"),
        ("Figures:", "6"),
        ("Supplementary files:",
         "S1 (PRISMA checklist), S2 (Search strategies), S3–S8 (data)"),
        ("Protocol registration:", "PROSPERO (CRD pending)"),
        ("Funding:", "[To be completed]"),
        ("Conflicts of interest:", "[To be completed]"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        r1 = p.add_run(label + " ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)


def build_highlights_page(doc):
    doc.add_heading("Highlights", level=1)
    for hl in HIGHLIGHTS:
        p = doc.add_paragraph(style="List Bullet")
        add_formatted_runs(p, hl)


def build_abstract_page(doc):
    doc.add_heading("Abstract", level=1)
    for label, body in ABSTRACT_BLOCKS:
        p = doc.add_paragraph()
        r = p.add_run(label + " ")
        r.bold = True
        add_formatted_runs(p, body)
    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    r.bold = True
    add_formatted_runs(p, KEYWORDS)


def build_toc_note(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "Table of Contents — to be auto-generated by Word "
        "(Insert > Table of Contents, or press F9 to update field).")
    r.italic = True
    r.font.size = Pt(10)


def build_supplementary_section(doc):
    hard_page_break(doc)
    doc.add_heading("Supplementary Material", level=1)
    for sid, desc in SUPPLEMENTARY_ITEMS:
        p = doc.add_paragraph()
        r = p.add_run(f"{sid}. ")
        r.bold = True
        add_formatted_runs(p, desc)


# ---------------------------------------------------------------------------
# Main render of body (Sections 1-5)
# ---------------------------------------------------------------------------

def render_body(doc, md_text, embedded_figs):
    lines = md_text.splitlines()
    # Find indices of Sections 1..5 ("## 1. ...") and "## References" / "## Supplementary"
    start_idx = None
    end_idx = None
    for i, ln in enumerate(lines):
        if re.match(r"^##\s+1\.\s", ln.strip()):
            start_idx = i
            break
    for i, ln in enumerate(lines):
        if ln.strip().startswith("## References"):
            end_idx = i
            break
    if start_idx is None or end_idx is None:
        raise RuntimeError("Couldn't find main text boundaries")

    body_lines = lines[start_idx:end_idx]

    i = 0
    n = len(body_lines)
    # Pre-compute lookups for figure / table anchors
    fig_anchors = FIGURE_ANCHORS
    tbl_anchors = TABLE_ANCHORS

    # We process line-by-line, but track current heading for anchor matching;
    # we emit figures/tables when we encounter the END of an anchored section
    # — defined as the line just before the next heading at the same or higher
    # level. Implementation: when we leave a subsection (next heading line)
    # we look back to its heading and emit any pending anchored assets.

    current_heading = None
    pending_anchor_emissions = []
    # Strategy: detect headings; whenever we hit a heading that should trigger
    # a figure/table, we set a flag and emit AFTER we've also emitted the
    # paragraphs of that section, i.e. just BEFORE the next heading line. We
    # track each anchor's emission state.

    pending = {"figs": [], "tables": []}

    def flush_pending():
        for fig_id in pending["figs"]:
            insert_figure(doc, fig_id, embedded_figs)
        for tbl_id in pending["tables"]:
            insert_table_file(doc, tbl_id)
        pending["figs"].clear()
        pending["tables"].clear()

    # Pre-strip the body by detecting markdown code blocks (PRISMA ASCII, etc.)
    in_code = False
    in_table = False
    table_buffer = []

    def emit_table_buffer():
        nonlocal table_buffer
        if not table_buffer:
            return
        # Only render inline tables that are NOT the three main tables
        # (those are inserted from tables/*.md files by anchor).
        # We render small inline summary tables verbatim.
        header, rows = parse_md_table(table_buffer)
        if header is not None:
            add_word_table(doc, header, rows, caption=None)
        table_buffer = []

    while i < n:
        line = body_lines[i]
        stripped = line.rstrip()

        # Code fence handling — skip ASCII PRISMA diagram blocks (we have
        # Figure 1 image instead).
        if stripped.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            i += 1
            continue

        # Heading detection — when we hit a heading, flush table buffer,
        # then flush any pending anchored figures/tables (they belong to the
        # previous section), then check whether this new heading triggers a
        # new pending emission, finally render the heading itself.
        m1 = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m1:
            emit_table_buffer()
            flush_pending()
            level_hashes = m1.group(1)
            heading_text = m1.group(2).strip()
            level = len(level_hashes) - 1  # ## -> level 1 (H1), ### -> H2
            level = max(1, min(level, 4))
            # Skip table/figure pseudo-headings embedded in markdown body
            if re.match(r"^Table\s+\d", heading_text) or re.match(
                    r"^Figure\s+\d", heading_text):
                # Skip — these are followed by markdown table or ASCII art
                # we handle separately. Advance i to consume; we do not emit
                # a heading.
                # We still need to capture data after these markers; just
                # skip the heading line.
                i += 1
                continue

            full_heading_line = stripped
            # Check anchors against the literal heading line
            for anchor_key, fig_list in fig_anchors.items():
                if full_heading_line == anchor_key:
                    pending["figs"].extend(fig_list)
            for anchor_key, tbl_list in tbl_anchors.items():
                if full_heading_line == anchor_key:
                    pending["tables"].extend(tbl_list)

            # Map markdown heading level to Word heading style
            word_level = level
            # Use level 1 for ## (sections), level 2 for ###, level 3 for ####
            doc.add_heading(heading_text, level=word_level)
            i += 1
            continue

        # Table block start
        if stripped.startswith("|"):
            table_buffer.append(stripped)
            i += 1
            continue
        else:
            if table_buffer:
                emit_table_buffer()

        # Blank line
        if not stripped.strip():
            i += 1
            continue

        # Horizontal rule
        if stripped.strip() == "---":
            i += 1
            continue

        # Bullet
        bullet_m = re.match(r"^([\-\*])\s+(.*)$", stripped)
        if bullet_m:
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, bullet_m.group(2))
            i += 1
            continue

        # Numbered list
        num_m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if num_m and len(num_m.group(1)) <= 2:
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, num_m.group(2))
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1

    # Final flushes
    emit_table_buffer()
    flush_pending()


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------

def main():
    md = load_manuscript()
    word_count = estimate_word_count(md)

    doc = Document()
    set_margins(doc, 1.0)
    configure_styles(doc)

    # Add page numbers in first section's footer
    add_page_number_footer(doc.sections[0])

    # Title page
    build_title_page(doc, word_count)
    hard_page_break(doc)

    # TOC note (Word will need user to insert; we leave a placeholder)
    build_toc_note(doc)
    hard_page_break(doc)

    # Highlights
    build_highlights_page(doc)
    hard_page_break(doc)

    # Abstract
    build_abstract_page(doc)
    hard_page_break(doc)

    # Main body Sections 1-5 with inline figures/tables
    embedded_figs = set()
    render_body(doc, md, embedded_figs)

    # References
    hard_page_break(doc)
    refs = parse_references(md)
    add_references_section(doc, refs)

    # Supplementary
    build_supplementary_section(doc)

    doc.save(str(OUT_DOCX))

    size = OUT_DOCX.stat().st_size
    print(f"WROTE {OUT_DOCX}")
    print(f"FILE_SIZE_BYTES={size}")
    print(f"FILE_SIZE_MB={size/1_048_576:.2f}")
    print(f"WORD_COUNT_MAIN_TEXT={word_count}")
    print(f"FIGURES_EMBEDDED={sorted(str(k) for k in embedded_figs)}")
    print(f"FIGURES_AVAILABLE={[k for k,v in FIGURES.items() if v.exists()]}")
    print(f"REFERENCES_PARSED={len(refs)}")


if __name__ == "__main__":
    main()
